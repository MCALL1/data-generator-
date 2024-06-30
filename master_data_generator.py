# -*- coding: utf-8 -*-
"""
Created on Mon January 16:24:17 2024

@author: mcall
"""

import pandas as pd
import numpy as np
from openpyxl import Workbook
from docx import Document

# Function to generate and save financial data
def generate_and_save_financial_data():
    num_rows = 1000000  # 1 million rows
    np.random.seed(42)
    quarters = np.arange(1, num_rows + 1)
    revenue = np.random.uniform(10000, 100000, size=num_rows)
    expenses = np.random.uniform(5000, 70000, size=num_rows)
    profit = revenue - expenses

    financial_data = pd.DataFrame({
        'Quarter': quarters,
        'Revenue': revenue,
        'Expenses': expenses,
        'Profit': profit
    })

    excel_file_path = "financial_data.xlsx"
    financial_data.to_excel(excel_file_path, index=False)
    print(f"Financial data saved to {excel_file_path}")

    doc = Document()
    doc.add_heading('Financial Data', level=1)
    doc.add_paragraph('This document contains financial data.')
    docx_file_path = "financial_data.docx"
    doc.save(docx_file_path)
    print(f"Financial data saved to {docx_file_path}")

# Function to generate and save product data
def generate_and_save_product_data():
    num_products = 1000000  # 1 million products
    np.random.seed(42)
    product_ids = np.arange(1, num_products + 1)
    product_names = ['Product_' + str(i) for i in range(1, num_products + 1)]
    prices = np.random.uniform(10, 1000, size=num_products)
    quantities = np.random.randint(1, 100, size=num_products)

    product_data = pd.DataFrame({
        'Product_ID': product_ids,
        'Product_Name': product_names,
        'Price': prices,
        'Quantity_Sold': quantities
    })

    excel_file_path = "product_data.xlsx"
    product_data.to_excel(excel_file_path, index=False)
    print(f"Product data saved to {excel_file_path}")

    doc = Document()
    doc.add_heading('Product Data', level=1)
    doc.add_paragraph('This document contains product data.')
    docx_file_path = "product_data.docx"
    doc.save(docx_file_path)
    print(f"Product data saved to {docx_file_path}")

# Main function to execute all tasks
def main():
    generate_and_save_financial_data()
    generate_and_save_product_data()

if __name__ == "__main__":
    main()
